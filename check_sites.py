import datetime
import socket
import ssl
import sys
from urllib.parse import urlparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_TABLE_ALIGNMENT # IMPORTATION AJOUTÉE POUR LE CENTRAGE
import requests


def set_cell_background(cell, fill_hex):
    """Applique une couleur de fond personnalisée à une cellule Word"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def check_website(url):
    if not url.startswith("http://") and not url.startswith("https://"):
        url = "https://" + url

    parsed_url = urlparse(url)
    domain = parsed_url.netloc
    est_https = url.startswith("https://")

    results = {
        "url": url,
        "disponible": "NON",
        "statut_http": "Inconnu",
        "last_modified": "Non fourni par le serveur",
        "certificat": "Inconnu",
        "jours_restants": 9999,
        "est_https": est_https,
        "erreur_ssl_critique": False,
        "statut_global": "DOWN",
    }

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "fr,fr-FR;q=0.8,en-US;q=0.5,en;q=0.3",
        "Cache-Control": "no-cache",
    }

    timeouts_config = (3.05, 12.0)

    try:
        response = requests.get(
            url, timeout=timeouts_config, allow_redirects=True, headers=headers
        )
        results["statut_http"] = str(response.status_code)
        code = response.status_code

        if code == 200:
            results["disponible"] = "OUI (Fonctionnel)"
            results["statut_global"] = "OK"
        elif 500 <= code <= 599:
            results["disponible"] = f"NON (Erreur Interne : {code})"
            results["statut_global"] = "DOWN"
        else:
            results["disponible"] = f"OUI (Statut Restreint : {code})"
            results["statut_global"] = "SPECIFIC"

        last_mod = response.headers.get("Last-Modified")
        if last_mod:
            results["last_modified"] = last_mod
        else:
            date_hdr = response.headers.get("Date")
            if date_hdr:
                results["last_modified"] = (
                    f"Inconnue (Généré le : {date_hdr})"
                )

    except requests.exceptions.SSLError:
        results["disponible"] = "OUI (Bloqué au niveau SSL)"
        results["statut_http"] = "Inconnu (SSL Error)"
        results["certificat"] = "INVALIDE / CORROMPU"
        results["erreur_ssl_critique"] = True
        results["statut_global"] = "DEGRADED"
        return results

    except (
        requests.exceptions.ConnectionError,
        requests.exceptions.Timeout,
        requests.exceptions.RequestException,
    ):
        results["disponible"] = "NON (Pas de réponse / Timeout / DNS)"
        results["statut_global"] = "DOWN"
        return results

    if est_https and not results["erreur_ssl_critique"]:
        try:
            context = ssl.create_default_context()
            with socket.create_connection((domain, 443), timeout=4) as sock:
                with context.wrap_socket(
                    sock, server_hostname=domain
                ) as ssock:
                    cert = ssock.getpeercert()
                    expire_str = cert.get("notAfter")
                    expire_date = datetime.datetime.strptime(
                        expire_str, "%b %d %H:%M:%S %Y %Z"
                    )
                    jours = (expire_date - datetime.datetime.utcnow()).days
                    results["jours_restants"] = jours

                    if jours > 0:
                        results["certificat"] = (
                            f"VALIDE ({jours} jours restants)"
                        )
                    else:
                        results["certificat"] = "EXPIRED (Expiré)"
                        results["erreur_ssl_critique"] = True
        except Exception as e:
            results["certificat"] = f"ERREUR TECHNIQUE ({type(e).__name__})"
            results["erreur_ssl_critique"] = True

    return results


def create_table_in_word(doc, title, data_list, header_color_hex, title_rgb):
    """Crée un tableau Word stylisé, centré et protégé contre les débordements"""
    full_title = f"{title}"
    heading = doc.add_heading(full_title, level=2)
    for run in heading.runs:
        run.font.color.rgb = title_rgb

    if not data_list:
        p = doc.add_paragraph()
        run = p.add_run("Aucun site détecté dans cette catégorie.")
        run.font.italic = True
        p.paragraph_format.space_after = Pt(12)
        return

    table = doc.add_table(rows=1, cols=6)
    table.autofit = False
    
    # OPTIMISATION 1 : Centrage horizontal parfait du tableau au milieu de la page
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    headers = [
        "N°",
        "Nom du site",
        "Disponible",
        "Code HTTP",
        "Dernière MAJ",
        "Certificat SSL",
    ]

    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], header_color_hex)
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = 1
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # Largeurs de colonnes ajustées au millimètre pour éviter TOUT débordement
    col_widths = [
        Inches(0.4),  # N°
        Inches(2.1),  # Nom du site (réduit pour éviter le débordement)
        Inches(1.3),  # Disponible
        Inches(0.7),  # Code HTTP
        Inches(1.4),  # Dernière MAJ
        Inches(1.4),  # Certificat SSL
    ]

    for index, item in enumerate(data_list, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        row_cells[1].text = str(item["url"])
        row_cells[2].text = str(item["disponible"])
        row_cells[3].text = str(item["statut_http"])
        row_cells[4].text = str(item["last_modified"])
        row_cells[5].text = str(item["certificat"])

        bg_color = "F5F5F5" if index % 2 == 0 else "FFFFFF"

        for cell in row_cells:
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                # Alignement à gauche pour les longues adresses, centré pour le reste
                for run in paragraph.runs:
                    run.font.size = Pt(9.0)  # Légèrement réduit pour maximiser l'espace interne

    # Appliquer strictement les largeurs de cellules et interdire le fractionnement
    for row in table.rows:
        # OPTIMISATION 2 : Empêcher les lignes de se briser bizarrement sur les sauts de page
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)


def process_file(filename):
    try:
        with open(filename, "r", encoding="utf-8") as file:
            sites = [line.strip() for line in file if line.strip()]
    except FileNotFoundError:
        print(f"Erreur : Le fichier '{filename}' n'existe pas dans ce dossier.")
        return

    tab_1_operationnel = []
    tab_2_degrade_ssl = []
    tab_3_non_securise = []
    tab_4_statut_specifique = []
    tab_5_indisponible_down = []
    tab_6_alerte_ssl_preventive = []

    print(f"--- Début de l'analyse : {len(sites)} plateformes chargées ---")
    
    for i, site in enumerate(sites, start=1):
        res = check_website(site)
        print(
            f"{i}. Analyse de : {res['url']} -> HTTP : {res['statut_http']} | Global : {res['statut_global']}"
        )

        if res["statut_global"] == "DOWN":
            tab_5_indisponible_down.append(res)
        elif res["erreur_ssl_critique"]:
            tab_2_degrade_ssl.append(res)
        elif not res["est_https"]:
            tab_3_non_securise.append(res)
        elif res["statut_global"] == "SPECIFIC":
            tab_4_statut_specifique.append(res)
        elif res["statut_global"] == "OK" and res["statut_http"] == "200":
            tab_1_operationnel.append(res)

        if (
            res["statut_global"] in ["OK", "SPECIFIC"]
            and res["jours_restants"] < 100
            and not res["erreur_ssl_critique"]
        ):
            tab_6_alerte_ssl_preventive.append(res)

    # =========================================================================
    # OPTIMISATION DE TRI AJOUTÉE : Ordre croissant sur la durée de validité SSL
    # =========================================================================
    tab_6_alerte_ssl_preventive.sort(key=lambda x: x["jours_restants"])
    # =========================================================================


    doc = Document()
    title_p = doc.add_heading(
        "Rapport d'Audit de Disponibilité et de Sécurité Web", 0
    )
    for run in title_p.runs:
        run.font.color.rgb = RGBColor(26, 54, 93)

    current_time_str = datetime.datetime.now().strftime("%d/%m/%Y à %H:%M")
    doc.add_paragraph(f"Généré le : {current_time_str}")

    # ===============================================
    # TABLEAU DE SYNTHÈSE VOLUMÉTRIQUE (CORRIGÉ ET STYLISÉ)
    # ==============================================
    
    doc.add_heading("Synthèse Volumétrique", level=2)

    p_total = doc.add_paragraph()
    run_t = p_total.add_run("Quantité total de sites audités : ")
    run_t.font.bold = True
    run_val = p_total.add_run(f"{len(sites)}\n")
    run_val.font.bold = True
    run_val.font.color.rgb = RGBColor(21, 101, 192)

    run_t2 = p_total.add_run("Nb: Un site peut se retrouver dans plusieurs catégories ")
    run_t2.font.bold = False
    run_t2.font.color.rgb = RGBColor(198, 40, 40)
    run_t2.font.size = Pt(9.5)  # <--- Modifiez cette valeur (ex: 9 ou 9.5) pour ajuster la taille


    # Création du tableau de résumé : 7 lignes (1 en-tête + 6 catégories), 2 colonnes
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.autofit = False

    # Configuration des en-têtes du tableau de synthèse
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = "Statut de la Plateforme"
    hdr_cells[1].text = "Quantité"

    # Style de l'en-tête (Fond Bleu Cobalt, Texte Blanc en Gras)
    for cell in hdr_cells:
        set_cell_background(cell, "1565C0")
        for paragraph in cell.paragraphs:
            paragraph.alignment = 0 if cell == hdr_cells[0] else 1  # Centrer la quantité
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # Préparation des données du tableau
    summary_data = [
        ("1. Opérationnels (200)", len(tab_1_operationnel), "FFFFFF"),
        ("2. Accessible mais Dégradé / Alerte SSL", len(tab_2_degrade_ssl), "FFFFFF"),
        ("3. Accessible mais Non sécurisé", len(tab_3_non_securise), "FFFFFF"),
        ("4. Statut Spécifique / Restrictions (Hors 200)", len(tab_4_statut_specifique), "FFFFFF"),
        #("5. Non disponible (Down)", len(tab_5_indisponible_down), "CRITIQUE"), # Tag spécial pour le rouge
        ("5. Non disponible (Down)", len(tab_5_indisponible_down), "FFFFFF"), 
        ("6. Alertes de Sécurité (Certificats SSL < 100 Jours)", len(tab_6_alerte_ssl_preventive), "FFFFFF")
    ]

    # Définition des largeurs des colonnes du résumé
    summary_widths = [Inches(4.5), Inches(1.5)]

    # Remplissage des lignes de données avec gestion des styles et du Zebra striping
    for idx, (label, count, status_type) in enumerate(summary_data, start=1):
        row_cells = summary_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(count)

        # Alternance de fond gris clair (Zebra striping) pour la lisibilité
        bg_color = "F9F9F9" if idx % 2 == 0 else "FFFFFF"

        for i, cell in enumerate(row_cells):
            set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                if i == 1:
                    paragraph.alignment = 1  # Centrer le chiffre de la quantité
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    # Si c'est la ligne Down et qu'il y a des pannes, on force l'écriture en rouge et gras
                    if status_type == "CRITIQUE" and count > 0:
                        run.font.color.rgb = RGBColor(198, 40, 40)
                        run.font.bold = True

    # Application des largeurs fixes sur toutes les cellules du tableau de synthèse
    for row in summary_table.rows:
        for i, width in enumerate(summary_widths):
            row.cells[i].width = width

    # Espace aéré avant d'attaquer les grands tableaux de détails
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(18)

    # =========================================================================
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    COLOR_OP_HEX, RGB_OP = "2E7D32", RGBColor(46, 125, 50)
    COLOR_DEG_HEX, RGB_DEG = "795548", RGBColor(121, 85, 72)
    COLOR_UNSEC_HEX, RGB_UNSEC = "607D8B", RGBColor(96, 125, 139)
    COLOR_SPEC_HEX, RGB_SPEC = "1565C0", RGBColor(21, 101, 192)
    COLOR_DOWN_HEX, RGB_DOWN = "C62828", RGBColor(198, 40, 40)
    COLOR_ALRT_HEX, RGB_ALRT = "EF6C00", RGBColor(239, 108, 0)


    # Génération séquentielle des 6 tableaux ordonnés
    create_table_in_word(
        doc,
        "1. Opérationnel (Status 200 + SSL valide)",
        tab_1_operationnel,
        COLOR_OP_HEX,RGB_OP,
        )
    
    create_table_in_word(doc,"2. Accessible mais Dégradé / Alerte SSL",tab_2_degrade_ssl,COLOR_DEG_HEX,RGB_DEG,)
    create_table_in_word(doc,"3. Accessible mais Non sécurisé",tab_3_non_securise,COLOR_UNSEC_HEX,RGB_UNSEC,)
    create_table_in_word(doc,"4. Statut Spécifique / Restrictions (Hors 200)",tab_4_statut_specifique,COLOR_SPEC_HEX,RGB_SPEC,)
    create_table_in_word(doc,"5. Non disponible (Down)",tab_5_indisponible_down,COLOR_DOWN_HEX,RGB_DOWN,)
    create_table_in_word(doc,"6. Alertes de Sécurité (Certificats SSL < 100 Jours)",tab_6_alerte_ssl_preventive,COLOR_ALRT_HEX,RGB_ALRT,)

    # 1. Récupérer la date et l'heure actuelles (Format : Jour-Mois-Année_Heures-Minutes)
    horodatage = datetime.datetime.now().strftime("%d-%m-%Y_%Hh-%M-%S")
    
    output_filename = f"Rapport_Site_Web_{horodatage}.docx"
    doc.save(output_filename)
    print(f"\n🎉 Rapport parfait généré avec succès : '{output_filename}'")


if __name__ == "__main__":
    # Nom du fichier par défaut ou passé en argument
    fichier_liste = sys.argv[1] if len(sys.argv) > 1 else "sites.txt"
    process_file(fichier_liste)